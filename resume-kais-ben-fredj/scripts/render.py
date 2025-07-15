import yaml
import os
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML
from docx import Document

# === Load Resume Data ===
def load_resume(file_path="resume.yaml"):
    with open(file_path, encoding="utf-8") as f:
        return yaml.safe_load(f)

# === HTML Render ===
def render_html(data, template_file="templates/html_template.jinja2", output_file="output/resume.html"):
    env = Environment(loader=FileSystemLoader("."))
    template = env.get_template(template_file)
    html_content = template.render(resume=data)
    
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"✅ HTML rendered: {output_file}")
    return output_file

# === PDF Render (via HTML) ===
def render_pdf(html_file="output/resume.html", output_file="output/resume.pdf"):
    HTML(html_file).write_pdf(output_file)
    print(f"✅ PDF generated: {output_file}")

# === Word Render ===
def render_word(data, output_file="output/resume.docx"):
    doc = Document()
    doc.add_heading(data["name"], 0)
    doc.add_paragraph(f"{data['contact']['email']} | {data['contact']['phone']} | {data['contact']['linkedin']}")
    
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(data["summary"])

    doc.add_heading("Technical Skills", level=1)
    for section, items in data["technical_skills"].items():
        doc.add_paragraph(f"{section.replace('_', ' ').title()}: {', '.join(items)}")

    doc.add_heading("Experience", level=1)
    for exp in data["experience"]:
        p = doc.add_paragraph()
        p.add_run(f"{exp['title']} – {exp['company']} ({exp['start']} to {exp['end']})\n").bold = True
        p.add_run(f"{exp['location']}\n")
        for bullet in exp["bullets"]:
            doc.add_paragraph(f"• {bullet}", style="List Bullet")

    doc.add_heading("Education", level=1)
    for edu in data["education"]:
        doc.add_paragraph(f"{edu['degree']}, {edu['institution']} ({edu['year']})")

    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    doc.save(output_file)
    print(f"✅ Word DOCX generated: {output_file}")

# === Run All ===
if __name__ == "__main__":
    resume_data = load_resume("resume.yaml")
    
    html_path = render_html(resume_data)
    render_pdf(html_path)
    render_word(resume_data)
